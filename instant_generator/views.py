from django.contrib.auth.decorators import login_required
from django.shortcuts import render, redirect
from django.contrib.sites.shortcuts import get_current_site
from django.contrib.auth import login, authenticate
from django.utils.encoding import force_bytes, force_text
from django.utils.http import urlsafe_base64_encode, urlsafe_base64_decode
from django.contrib import messages
from django.template.loader import render_to_string
from .models import User, InstantGenerator
from .forms import UserForm, ProfileForm, InstantGeneratorForm
from .tokens import account_activation_token
from django.db import transaction

from io import BytesIO
from reportlab.pdfgen import canvas
from django.http import HttpResponse, response

from docx import Document
from docx.shared import Inches


from reportlab.platypus import SimpleDocTemplate, Paragraph, PageBreak, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import mm, inch


# Create your views here.
def signup(request):
    if request.method == 'POST':
        form = UserForm(request.POST)
        if form.is_valid():
            user = form.save()
            user.is_active = True
            user.save()
            current_site = get_current_site(request)
            subject = 'Activate Your Tool-X Account'
            message = render_to_string('registration/activation_email.html', {
                'user': user,
                'domain': current_site.domain,
                'uid': urlsafe_base64_encode(force_bytes(user.pk)),
                'token': account_activation_token.make_token(user),
            })
            user.email_user(subject, message)
            return redirect('activation_sent')

    else:
        form = UserForm()

    return render(request, 'registration/signup.html', {'form': form})


def activation_sent(request):

    return render(request, 'registration/activation_sent.html', {})


def activate(request, uidb64, token):
    try:
        uid = urlsafe_base64_decode(uidb64).decode()
        user = User.objects.get(pk=uid)
    except (TypeError, ValueError, OverflowError, User.DoesNotExist):
        user = None

    if user is not None and account_activation_token.check_token(user, token):
        user.is_active = True
        user.email_confirmed = True
        user.save()
        login(request, user, backend='django.contrib.auth.backends.ModelBackend')
        return redirect('index')
    else:
        return render(request, 'registration/activation_invalid.html')


def profile(request):

    return render(request, 'registration/profile.html', {'user': request.user})


@transaction.atomic
def edit_profile(request):
    if request.method == 'POST':
        user_form = UserForm(request.POST, instance=request.user)
        form = ProfileForm(request.POST, instance=request.user)
        if user_form.is_valid() and form.is_valid():
            user_form.save()
            form.save()
            messages.success(request, 'Your profile was successfully updated!')
            return redirect('profile')

        else:
            messages.error(request, 'Please correct the error below.')

    else:
        user_form = UserForm(instance=request.user)
        form = ProfileForm(instance=request.user)

    return render(request, 'registration/edit_profile.html', {
        'user_form': user_form,
        'form': form
    })


def dashboard(request):

    return render(request, 'instant_generator/dashboard.html', {})


@login_required
@transaction.atomic
def create(request):
    if request.method == 'POST':
        form = InstantGeneratorForm(request.POST)
        if form.is_valid():
            instance = form.save(commit=False)
            instance.user = request.user
            instance.save()
            return redirect('congratulation')

    else:
        form = InstantGeneratorForm()

    return render(request, 'instant_generator/create.html', {'form': form})


def congratulation(request):

    return render(request, 'instant_generator/congratulation.html', {})


def my_adcopies(request):
    # adcopies = InstantGenerator.objects.all().order_by('-created_on')
    adcopies = InstantGenerator.objects.filter(user=request.user)
    context = {
        "adcopies": adcopies,
    }

    return render(request, 'instant_generator/my_adcopies.html', context)


def preview(request, pk):
    generated = InstantGenerator.objects.get(pk=pk)
    context = {
        "generated": generated,
    }

    return render(request, 'instant_generator/preview.html', context)


def pdf(request, pk):
    generated = InstantGenerator.objects.get(user=request.user, pk=pk)
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="Sales Letter.pdf"'

    pdf_buffer = BytesIO()
    # pagesize = (140 * mm, 216 * mm)  # width, height
    my_doc = SimpleDocTemplate(
        pdf_buffer,
        # pagesize=pagesize
    )

    flowables = []
    my_doc.build(flowables)
    sample_style_sheet = getSampleStyleSheet()
    # sample_style_sheet.list()

    paragraph_1 = Paragraph(generated.Get_Attention, sample_style_sheet['Title'])
    paragraph_2 = Paragraph(
        generated.Identify_the_Problem_Your_Audience_Have,
        sample_style_sheet['BodyText']
    )
    paragraph_3 = Paragraph(
        generated.Provide_the_Solution,
        sample_style_sheet['BodyText']
    )
    paragraph_4 = Paragraph(
        generated.Present_your_Credentials,
        sample_style_sheet['BodyText']
    )
    paragraph_5 = Paragraph(
        generated.Show_the_Benefits,
        sample_style_sheet['BodyText']
    )
    paragraph_6 = Paragraph(
        generated.Give_Social_Proof,
        sample_style_sheet['BodyText']
    )
    paragraph_7 = Paragraph(
        generated.Make_Your_Offer,
        sample_style_sheet['Heading4']
    )
    paragraph_8 = Paragraph(
        generated.Give_a_Guarantee,
        sample_style_sheet['Italic']
    )
    paragraph_9 = Paragraph(
        generated.Inject_Scarcity,
        sample_style_sheet['BodyText']
    )
    paragraph_10 = Paragraph(
        generated.Call_to_action,
        sample_style_sheet['Heading3']
    )
    paragraph_11 = Paragraph(
        generated.Give_a_Warning,
        sample_style_sheet['BodyText']
    )
    paragraph_12 = Paragraph(
        generated.Close_with_a_Reminder,
        sample_style_sheet['Heading4']
    )

    flowables.append(paragraph_1)
    flowables.append(paragraph_2)
    flowables.append(paragraph_3)
    flowables.append(paragraph_4)
    flowables.append(paragraph_5)
    flowables.append(paragraph_6)
    flowables.append(paragraph_7)
    flowables.append(paragraph_8)
    flowables.append(paragraph_9)
    flowables.append(paragraph_10)
    flowables.append(paragraph_11)
    flowables.append(paragraph_12)

    my_doc.build(flowables)

    pdf = pdf_buffer.getvalue()
    pdf_buffer.close()
    response.write(pdf)

    return response


def docx(request, pk):
    generated = InstantGenerator.objects.get(user=request.user, pk=pk)
    document = Document()
    document.add_heading(generated.Get_Attention, 0)
    document.add_paragraph(generated.Identify_the_Problem_Your_Audience_Have)
    document.add_paragraph(generated.Provide_the_Solution)
    document.add_paragraph(generated.Present_your_Credentials)
    document.add_paragraph(generated.Show_the_Benefits)
    document.add_paragraph(generated.Give_Social_Proof)
    document.add_paragraph(generated.Make_Your_Offer, style='IntenseQuote')
    document.add_paragraph(generated.Give_a_Guarantee)
    document.add_paragraph(generated.Inject_Scarcity)
    document.add_paragraph(generated.Call_to_action)
    document.add_paragraph(generated.Give_a_Warning)

    p = document.add_paragraph()
    p.add_run(generated.Close_with_a_Reminder).bold = True

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=Sales Letter.docx'
    document.save(response)

    return response

